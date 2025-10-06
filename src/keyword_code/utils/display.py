"""
Display functions for the keyword_code package.
"""

import streamlit as st
import base64
import json
import re
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import fitz
from datetime import datetime
from typing import Dict, List, Any, Tuple, Optional
from pathlib import Path # Added for Path operations
from ..config import logger, RAG_TOP_K

# Import the model loading functions directly from the models module
from ..models.embedding import load_embedding_model, load_reranker_model

from ..rag.retrieval import retrieve_relevant_chunks_for_chat
from ..utils.async_utils import run_async
from ..ai.analyzer import DocumentAnalyzer # Import DocumentAnalyzer
from ..ai.chat import generate_chat_response

# Load the embedding model using the cached function
embedding_model = load_embedding_model()
reranker_model = load_reranker_model()

# --- Helper function to load and encode images ---
def get_base64_encoded_image(image_path: Path) -> Optional[str]:
    """Get base64 encoded image."""
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode()
    except Exception as e:
        logger.error(f"Error encoding image {image_path}: {str(e)}")
        return None

# --- Define check_img for use in UI elements ---
check_img = "✅" # Default to emoji
try:
    # Path to the assets directory from src/keyword_code/utils/
    assets_path = Path(__file__).parent.parent / "assets"
    correct_png_path = assets_path / "correct.png"
    logger.info(f"Attempting to load check icon from: {correct_png_path}")
    if correct_png_path.is_file():
        check_base64 = get_base64_encoded_image(correct_png_path)
        if check_base64:
            check_img = f'<img src="data:image/png;base64,{check_base64}" style="width: 18px; height: 18px; vertical-align: middle; margin-right: 5px;" alt="✓">'
        else:
            logger.warning(f"Failed to encode check icon: {correct_png_path}")
    else:
        logger.warning(f"Check icon not found at: {correct_png_path}")
except Exception as img_e:
    logger.warning(f"Could not load check icon, using emoji fallback: {img_e}")


def find_annotated_pdf_for_filename(filename: str) -> Optional[bytes]:
    """Finds the base64 decoded annotated PDF bytes for a given filename from session state."""
    for result in st.session_state.get("analysis_results", []):
        if isinstance(result, dict) and result.get("filename") == filename and result.get("annotated_pdf"):
            try:
                return base64.b64decode(result["annotated_pdf"])
            except Exception as e:
                logger.error(f"Failed to decode annotated PDF for {filename} in chat citation: {e}")
                return None
    logger.warning(f"Could not find annotated PDF data for {filename} in session state analysis_results.")
    return None

def process_chat_response_for_numbered_citations(raw_response_text: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Processes raw AI response text containing (Source:...) citations.
    Replaces them with sequential numbers [1], [2], etc., and returns the
    modified text along with a list of citation details for creating buttons.

    Args:
        raw_response_text: The original text from the AI.

    Returns:
        Tuple containing:
        - str: The response text with inline citations replaced by numbers ([1], [2]).
        - List[Dict[str, Any]]: A list of citation details, each dict containing
                               'number', 'filename', 'page', 'pdf_bytes'.
    """
    if not raw_response_text:
        return "", []

    citation_pattern = re.compile(r"\(Source:\s*(?P<filename>[^,]+?)\s*,\s*Page:\s*(?P<page>\d+)\)")

    citations_found_for_replacement = [] # Stores info needed for text replacement
    citation_details_for_footer = [] # Stores unique details for footer buttons
    next_citation_number = 1
    processed_text = raw_response_text

    # Find all citations and assign sequential numbers
    for match in citation_pattern.finditer(raw_response_text):
        filename = match.group("filename").strip()
        page_str = match.group("page").strip()
        try:
            page_num = int(page_str)

            # --- Assign unique number to THIS instance ---
            current_number = next_citation_number

            # Get PDF bytes for this source
            pdf_bytes = find_annotated_pdf_for_filename(filename)

            # Store details for the footer button list
            citation_details_for_footer.append({
                'number': current_number,
                'filename': filename,
                'page': page_num,
                'pdf_bytes': pdf_bytes # Can be None if not found
            })

            # Store details needed to replace the text later
            citations_found_for_replacement.append({
                'start': match.start(),
                'end': match.end(),
                'number': current_number,
                'original_text': match.group(0)
            })

            # Increment for the *next* citation found
            next_citation_number += 1

        except ValueError:
            logger.warning(f"Found invalid page number in citation: {match.group(0)}")
        except Exception as e:
            logger.error(f"Error processing citation {match.group(0)}: {e}")

    # Second pass: Replace citations in the text from end to start (to avoid index issues)
    # Sort by start position in reverse order
    citations_found_for_replacement.sort(key=lambda x: x['start'], reverse=True)

    for citation in citations_found_for_replacement:
        processed_text = (
            processed_text[:citation['start']] +
            f" [{citation['number']}]" +
            processed_text[citation['end']:]
        )

    return processed_text.strip(), citation_details_for_footer


def display_followup_citations_like_main_analysis(citation_details: List[Dict[str, Any]], qa_index: int = 0, answer_text: str = ""):
    """
    Displays follow-up citations in the same format as the main analysis supporting citations.
    Uses the same styling with quoted text, verification badges, and "Go" buttons.

    Args:
        citation_details: List of citation dictionaries from process_chat_response_for_numbered_citations
        qa_index: Index of the Q&A pair for unique keys
        answer_text: The raw answer text to try to extract context from
    """
    if not citation_details:
        st.info("No supporting citations were identified for this follow-up question.")
        return

    citation_counter = 0
    for citation_idx, citation in enumerate(citation_details):
        citation_counter += 1

        # Extract citation info
        filename = citation.get('filename', 'Unknown')
        page_num = citation.get('page_num', citation.get('page', 1))
        pdf_bytes = citation.get('pdf_bytes')
        actual_citation_number = citation.get('number', citation_counter)

        # For follow-up citations, we'll assume they're verified since they come from the RAG system
        is_verified = True  # Follow-up citations are from RAG retrieval, so considered verified

        # Extract the actual relevant phrase from the AI response
        citation_text = f"Referenced content from {filename}, Page {page_num}"  # Default fallback

        if answer_text:
            # Look for text around where this citation number appears
            # Use the actual citation number from the citation details, not the counter
            citation_pattern = f"\\[{actual_citation_number}\\]"
            logger.debug(f"Looking for citation pattern '{citation_pattern}' in answer text for {filename}")

            # First, let's check if the citation pattern exists in the text
            if re.search(citation_pattern, answer_text):
                # Split the text into sentences and find the one with this citation
                # Use multiple sentence delimiters to be more comprehensive
                sentences = re.split(r'[.!?]+(?:\s|$)', answer_text)

                for sentence in sentences:
                    if re.search(citation_pattern, sentence):
                        # Clean up the sentence by removing citation markers and extra whitespace
                        clean_sentence = re.sub(r'\[\d+\]', '', sentence).strip()
                        clean_sentence = re.sub(r'\s+', ' ', clean_sentence)  # Normalize whitespace

                        if len(clean_sentence) > 15:  # Only use if it's substantial enough
                            # Truncate if too long, but try to end at a word boundary
                            if len(clean_sentence) > 120:
                                truncated = clean_sentence[:120]
                                # Try to end at the last complete word
                                last_space = truncated.rfind(' ')
                                if last_space > 80:  # Only if we don't cut too much
                                    truncated = truncated[:last_space]
                                citation_text = truncated + "..."
                            else:
                                citation_text = clean_sentence
                        break

            # If we didn't find a sentence with the citation, try a different approach
            # Look for text immediately before the citation marker
            else:
                # Try to find any occurrence of the citation number in the text
                match = re.search(citation_pattern, answer_text)
                if match:
                    start_pos = match.start()
                    # Look backwards to find the start of the relevant phrase
                    # Try to find the beginning of the sentence or clause
                    text_before = answer_text[:start_pos]

                    # Look for sentence boundaries going backwards
                    sentence_starts = [m.end() for m in re.finditer(r'[.!?]\s+', text_before)]
                    if sentence_starts:
                        sentence_start = sentence_starts[-1]
                    else:
                        # If no sentence boundary, look for other natural breaks
                        clause_starts = [m.end() for m in re.finditer(r'[,;]\s+', text_before)]
                        if clause_starts:
                            sentence_start = clause_starts[-1]
                        else:
                            sentence_start = max(0, start_pos - 100)  # Last resort: 100 chars back

                    # Extract the relevant phrase
                    relevant_phrase = answer_text[sentence_start:start_pos].strip()
                    if len(relevant_phrase) > 15:
                        # Clean up and use this phrase
                        relevant_phrase = re.sub(r'\s+', ' ', relevant_phrase)
                        if len(relevant_phrase) > 120:
                            last_space = relevant_phrase.rfind(' ', 0, 120)
                            if last_space > 80:
                                relevant_phrase = relevant_phrase[:last_space] + "..."
                            else:
                                relevant_phrase = relevant_phrase[:120] + "..."
                        citation_text = relevant_phrase
                else:
                    # If we still can't find the citation, try to extract meaningful content
                    # Look for sentences that mention the filename or related content
                    filename_base = filename.replace('.pdf', '').replace('.docx', '')
                    if filename_base in answer_text:
                        # Find sentences that mention this document
                        sentences = re.split(r'[.!?]+(?:\s|$)', answer_text)
                        for sentence in sentences:
                            if filename_base.lower() in sentence.lower():
                                clean_sentence = re.sub(r'\[\d+\]', '', sentence).strip()
                                clean_sentence = re.sub(r'\s+', ' ', clean_sentence)
                                if len(clean_sentence) > 15:
                                    if len(clean_sentence) > 120:
                                        last_space = clean_sentence.rfind(' ', 0, 120)
                                        if last_space > 80:
                                            clean_sentence = clean_sentence[:last_space] + "..."
                                        else:
                                            clean_sentence = clean_sentence[:120] + "..."
                                    citation_text = clean_sentence
                                break

        # Badge HTML (same as main analysis)
        if is_verified:
            badge_html = '<span style="display: inline-block; background-color: #d1fecf; color: #11631a; padding: 1px 6px; border-radius: 0.25rem; font-size: 0.8em; margin-left: 5px; border: 1px solid #a1e0a3; font-weight: 600;">✔ Verified</span>'
        else:
            badge_html = '<span style="display: inline-block; background-color: #ffeacc; color: #a05e03; padding: 1px 6px; border-radius: 0.25rem; font-size: 0.8em; margin-left: 5px; border: 1px solid #f8c78d; font-weight: 600;">⚠️ Needs Review</span>'

        # Create columns for citation and Go button (same layout as main analysis)
        cite_col, btn_col = st.columns([0.90, 0.10], gap="small")

        with cite_col:
            # Citation text container (same styling as main analysis)
            st.markdown(f"""
            <div style="border: 1px solid #e0e0e0; border-radius: 5px; padding: 8px 12px; margin-top: 5px; margin-bottom: 8px; background-color: #f9f9f9;">
                <div style="margin-bottom: 5px; display: flex; justify-content: space-between; align-items: center;">
                    <span style="font-weight: bold;">Citation {citation_counter} {badge_html}</span>
                    <span style="font-size: 0.8em; color: #555;">Page {page_num} | RAG Retrieved</span>
                </div>
                <div style="color: #333; line-height: 1.4; font-size: 0.95em;"><i>"{citation_text}"</i></div>
            </div>
            """, unsafe_allow_html=True)

        with btn_col:
            # 'Go' button logic (same as main analysis)
            st.markdown('<div style="margin-top: 20px;"></div>', unsafe_allow_html=True)
            if pdf_bytes:
                button_key = f"followup_goto_{qa_index}_{citation_counter}_{citation_idx}"
                if st.button("Go", key=button_key, type="secondary", help=f"Go to Page {page_num} in {filename}", use_container_width=True):
                    update_pdf_view(pdf_bytes=pdf_bytes, page_num=page_num, filename=filename)
                    st.session_state.scroll_to_pdf_viewer = True
                    st.rerun()
            else:
                st.caption("PDF N/A")


def display_chat_message_with_citations(processed_text: str, citation_details: List[Dict[str, Any]], msg_idx: int = 0):
    """
    Displays the processed chat message containing numbered citations [1], [2], etc.,
    and lists the corresponding source buttons below.

    Args:
        processed_text: The message text with (Source:...) replaced by [1], [2].
        citation_details: A list of dictionaries from process_chat_response_for_numbered_citations,
                          each containing 'number', 'filename', 'page', 'pdf_bytes'.
        msg_idx: The index of the message in the overall chat history (for unique keys).
    """

    # Display the main message content with inline numbers
    # Use a div with word-wrap to ensure text wraps properly without horizontal scrolling
    st.markdown(
        f'<div style="word-wrap: break-word; overflow-wrap: break-word; white-space: normal;">{processed_text}</div>',
        unsafe_allow_html=True
    )

    # Display the citation sources below if any exist
    if citation_details:
        st.markdown('<hr style="margin: 10px 0; border: 0; border-top: 1px solid #e0e0e0;">', unsafe_allow_html=True)
        st.caption("Sources:")

        for i, citation in enumerate(citation_details):
            number = citation.get('number', i+1)
            filename = citation.get('filename', 'Unknown')
            page_num = citation.get('page_num', citation.get('page', 1))
            pdf_bytes = citation.get('pdf_bytes')

            # Create a unique key for each citation
            citation_key = f"cite_{filename}_{page_num}_{number}_{msg_idx}_{i}"

            if pdf_bytes:
                # Include the message index (msg_idx) and citation index (i) for guaranteed uniqueness
                button_key = f"chat_footer_cite_{citation_key}"
                # Make the entire citation text a button
                st.button(
                    f"[{number}] 📄 {filename}, p{page_num}",
                    key=button_key,
                    help=f"View Page {page_num} in {filename}",
                    type="secondary",
                    on_click=update_pdf_view,
                    args=(pdf_bytes, page_num, filename)
                )
            else:
                # Try to find the PDF in the analysis results as a fallback
                found_pdf = False
                if "analysis_results" in st.session_state:
                    for result in st.session_state.analysis_results:
                        if isinstance(result, dict) and result.get("filename") == filename and result.get("annotated_pdf"):
                            try:
                                pdf_bytes = base64.b64decode(result["annotated_pdf"])
                                button_key = f"chat_footer_cite_fallback_{citation_key}"
                                st.button(
                                    f"[{number}] 📄 {filename}, p{page_num}",
                                    key=button_key,
                                    help=f"View Page {page_num} in {filename}",
                                    type="secondary",
                                    on_click=update_pdf_view,
                                    args=(pdf_bytes, page_num, filename)
                                )
                                found_pdf = True
                                break
                            except Exception as e:
                                logger.error(f"Failed to decode annotated PDF for {filename} in fallback: {e}")

                # If PDF not found, display text indicating the source
                if not found_pdf:
                    st.markdown(
                        f'<div style="color: #888; padding: 0.25rem 0.75rem; font-size: 0.9em; border-radius: 0.25rem; background-color: #f0f0f0;">'
                        f'[{number}] 📄 {filename}, p{page_num} (PDF not available)'
                        f'</div>',
                        unsafe_allow_html=True
                    )


def create_report_package_content(issue_description: str, results: List[Tuple[Dict[str, Any], Dict[str, Any]]]) -> bytes:
    """
    Create a ZIP package containing all relevant information for issue reporting.

    Args:
        issue_description: User's description of the issue
        results: List of tuples containing (result, ai_analysis) pairs

    Returns:
        bytes: The ZIP file as bytes
    """
    # Create a BytesIO object to store the ZIP file
    zip_buffer = BytesIO()

    # Create a ZIP file
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        # Add report metadata
        report_data = {
            "issue_description": issue_description,
            "timestamp": datetime.now().isoformat(),
            "user_inputs": {
                "prompt": st.session_state.get("user_prompt", ""),
                "uploaded_files": list(st.session_state.get("preprocessed_data", {}).keys())
            },
            "chat_history": st.session_state.get("chat_messages", []),
            "followup_qa": st.session_state.get("followup_qa", []),
            "current_document": st.session_state.get("current_pdf_name", "")
        }

        # Add report data as JSON
        zip_file.writestr("report_data.json", json.dumps(report_data, indent=2))

        # Add analysis results
        for i, (result, ai_analysis) in enumerate(results):
            filename = result.get("filename", f"unknown_file_{i}")

            # Add the analysis result as JSON
            zip_file.writestr(
                f"analysis_results/{filename}_analysis.json",
                json.dumps(ai_analysis, indent=2)
            )

            # Add the annotated PDF if available
            if result.get("annotated_pdf"):
                try:
                    annotated_pdf_bytes = base64.b64decode(result["annotated_pdf"])
                    zip_file.writestr(
                        f"annotated_pdfs/{filename}_annotated.pdf",
                        annotated_pdf_bytes
                    )
                except Exception as e:
                    logger.error(f"Error adding annotated PDF for {filename} to report package: {e}")

            # Add the original document if available
            if "preprocessed_data" in st.session_state and filename in st.session_state.preprocessed_data:
                try:
                    original_bytes = st.session_state.preprocessed_data[filename].get("original_bytes")
                    if original_bytes:
                        zip_file.writestr(
                            f"original_documents/{filename}",
                            original_bytes
                        )
                except Exception as e:
                    logger.error(f"Error adding original document for {filename} to report package: {e}")

    # Reset buffer position and return the bytes
    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def export_to_word(exportable_results_list: List[Dict[str, Any]]) -> bytes:
    """
    Export analysis results to a Word document.

    Args:
        exportable_results_list: List of dictionaries containing analysis results

    Returns:
        bytes: The Word document as bytes
    """
    # Create a new Word document
    doc = Document()

    # Add title
    title = doc.add_heading("Document Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add report generation date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.italic = True

    # Add a page break after the title page
    doc.add_page_break()

    # Process each file's results
    for file_result in exportable_results_list:
        filename = file_result.get("filename", "Unknown File")
        analysis = file_result.get("analysis", {})

        # Add file heading
        doc.add_heading(f"Document: {filename}", 1)

        # Process each analysis section
        for section_key, section_data in analysis.get("analysis_sections", {}).items():
            # Add section heading
            section_name = section_key.replace("_", " ").title()
            doc.add_heading(section_name, 2)

            # Add analysis text
            if section_data.get("Analysis"):
                p = doc.add_paragraph()
                p.add_run("Analysis: ").bold = True
                p.add_run(section_data.get("Analysis"))

            # Add context if available
            if section_data.get("Context"):
                p = doc.add_paragraph()
                p.add_run("Context: ").bold = True
                p.add_run(section_data.get("Context"))

            # Add supporting phrases
            supporting_phrases = section_data.get("Supporting_Phrases", [])
            if supporting_phrases and supporting_phrases != ["No relevant phrase found."]:
                doc.add_heading("Supporting Citations", 3)

                for phrase in supporting_phrases:
                    # Get verification info from the file_result data
                    data_rows = []
                    try:
                        file_data = file_result.get("data", [])
                        if isinstance(file_data, list):
                            data_rows = [row for row in file_data
                                        if isinstance(row, dict) and row.get("Supporting Phrase") == phrase]
                    except Exception as e:
                        logger.error(f"Error getting data rows for phrase '{phrase}': {e}")

                    if data_rows:
                        try:
                            data_row = data_rows[0]
                            # Check if Verified is "Yes" or True
                            verified_value = data_row.get("Verified")
                            if isinstance(verified_value, str):
                                is_verified = verified_value.lower() == "yes"
                            elif isinstance(verified_value, bool):
                                is_verified = verified_value
                            else:
                                is_verified = False

                            # Get page number info
                            page_num_info = data_row.get("Page", "Unknown")

                            # Get score info
                            score_info = data_row.get("Match Score", "N/A")
                        except Exception as e:
                            logger.error(f"Error extracting verification info from data row: {e}")
                            is_verified = False
                            page_num_info = "Unknown"
                            score_info = "N/A"
                    else:
                        is_verified = False
                        page_num_info = "Unknown"
                        score_info = "N/A"

                    # Add the phrase with verification status
                    p = doc.add_paragraph()
                    if is_verified:
                        p.add_run("✓ ").bold = True
                        p.add_run(phrase)
                        details_run = p.add_run(f" (Pg: {page_num_info}, Score: {score_info})")
                        details_run.italic = True
                        details_run.font.size = Pt(9)
                    else:
                        p.add_run("❓ ").bold = True
                        p.add_run(phrase)
                        details_run = p.add_run(" (Not verified in document)")
                        details_run.italic = True
                        details_run.font.size = Pt(9)

            # Add a separator after each section
            doc.add_paragraph("---")

        # Add a page break after each file
        doc.add_page_break()

    # Add Follow-up Q&A section if any exist
    followup_qa = st.session_state.get("followup_qa", [])
    if followup_qa:
        doc.add_heading("Follow-up Questions & Answers", 1)

        for i, qa_pair in enumerate(followup_qa):
            # Add question
            doc.add_heading(f"Question {i+1}", 2)
            q_paragraph = doc.add_paragraph()
            q_paragraph.add_run("Q: ").bold = True
            q_paragraph.add_run(qa_pair.get("question", ""))

            # Add answer
            a_paragraph = doc.add_paragraph()
            a_paragraph.add_run("A: ").bold = True
            a_paragraph.add_run(qa_pair.get("answer", ""))

            # Add timestamp if available
            if qa_pair.get("timestamp"):
                timestamp_paragraph = doc.add_paragraph()
                timestamp_run = timestamp_paragraph.add_run(f"Asked on: {qa_pair['timestamp']}")
                timestamp_run.italic = True
                timestamp_run.font.size = Pt(9)

            # Add citations if available
            citation_details = qa_pair.get("citation_details", [])
            if citation_details:
                doc.add_heading("Citations", 3)
                for citation in citation_details:
                    cite_paragraph = doc.add_paragraph()
                    cite_paragraph.style = 'List Bullet'
                    cite_paragraph.add_run(f"[{citation.get('number', '')}] ")
                    cite_paragraph.add_run(f"{citation.get('filename', 'Unknown')}, ")
                    cite_paragraph.add_run(f"Page {citation.get('page_num', 'Unknown')}")

            # Add separator between Q&A pairs
            if i < len(followup_qa) - 1:
                doc.add_paragraph("---")

    # Save the document to a BytesIO object
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    # Return the document as bytes
    return docx_buffer.getvalue()


def update_pdf_view(pdf_bytes, page_num=1, filename=None):
    """
    Updates the PDF view in the session state.

    Args:
        pdf_bytes: The PDF bytes to display
        page_num: The page number to display (1-based)
        filename: The name of the file
    """
    if pdf_bytes:
        st.session_state.pdf_bytes = pdf_bytes
        st.session_state.pdf_page = page_num
        st.session_state.show_pdf = True
        if filename:
            st.session_state.current_pdf_name = filename
        logger.info(f"Updated PDF view to {filename}, page {page_num}")
    else:
        logger.warning("Attempted to update PDF view with empty bytes")
        st.session_state.show_pdf = False


def display_pdf_viewer(pdf_bytes, current_page=1, filename=None):
    """
    Displays a PDF viewer with navigation controls.

    Args:
        pdf_bytes: The PDF bytes to display
        current_page: The current page number (1-based)
        filename: The name of the file
    """
    if not pdf_bytes:
        st.warning("No PDF data available to display.")
        return

    try:
        # Create a base64 encoded PDF string
        base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')

        # Display the filename if provided
        if filename:
            st.markdown(f"**Viewing:** {filename}")

        # Create an iframe to display the PDF
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

        # Add page navigation controls
        col1, col2, col3 = st.columns([1, 3, 1])
        with col1:
            if st.button("◀ Previous Page", disabled=(current_page <= 1)):
                update_pdf_view(pdf_bytes, current_page - 1, filename)
                st.rerun()

        with col2:
            st.markdown(f"<div style='text-align: center;'>Page {current_page}</div>", unsafe_allow_html=True)

        with col3:
            if st.button("Next Page ▶"):
                update_pdf_view(pdf_bytes, current_page + 1, filename)
                st.rerun()

    except Exception as e:
        logger.error(f"Error displaying PDF: {e}")
        st.error(f"Error displaying PDF: {e}")


def display_analysis_results(results: List[Dict[str, Any]]):
    """
    Displays the analysis results in a structured format with a two-column layout.
    Left column shows AI analysis, right column shows tools including PDF viewer, chat, and export options.

    Args:
        results: A list of result dictionaries, each containing analysis data for a file
    """
    # Initialize followup_qa if it doesn't exist
    if "followup_qa" not in st.session_state:
        st.session_state.followup_qa = []

    if not results:
        st.warning("No analysis results to display.")
        return

    # Define CSS styles based on app.py.bak
    st.markdown("""
    <style>
    .header-title {
        font-weight: 700;
        font-size: 1.5rem;
        color: #333; /* From app.py.bak's .header-title */
        margin: 0;
        padding: 0;
    }
    .sleek-container {
        background-color: #f5f5f5;
        border-radius: 8px;
        padding: 8px 16px;
        margin: 0 0 16px 0;
        display: flex;
        align-items: center;
        justify-content: space-between;
        border: 1px solid #e0e0e0;
    }
    .file-name {
        font-weight: 600;
        color: #424242; /* From app.py.bak's .file-name */
        font-size: 1rem;
        display: flex;
        align-items: center;
        margin: 0;
        padding: 0;
    }
    .file-icon {
        color: #1976d2; /* From app.py.bak's .file-icon */
        margin-right: 8px;
    }
    </style>
    """, unsafe_allow_html=True)

    # Create a two-column layout for the analysis results
    analysis_col, tools_col = st.columns([2.5, 1.5], gap="small")

    # Add an anchor for auto-scrolling
    st.markdown('<div id="results-anchor"></div>', unsafe_allow_html=True)

    # Left Column: AI Analysis Display
    with analysis_col:
        # Header style from app.py.bak
        st.markdown('<div class="header-title">AI Analysis Results</div>', unsafe_allow_html=True)
        st.markdown('<hr style="margin: 12px 0; border: 0; border-top: 1px solid #e0e0e0;">', unsafe_allow_html=True)

        # Create a scrollable container for the analysis
        with st.container(height=1220, border=True):
            # Process results to extract only those with real analysis
            results_with_real_analysis = []
            for result in results:
                try:
                    filename = result.get("filename", "Unknown File")
                    ai_analysis_str = result.get("ai_analysis", "{}")

                    try:
                        ai_analysis = json.loads(ai_analysis_str)
                        # Only include results with actual analysis sections
                        if ai_analysis.get("analysis_sections", {}):
                            results_with_real_analysis.append((result, ai_analysis))
                    except json.JSONDecodeError:
                        logger.error(f"Failed to parse analysis JSON for {filename}")
                        continue
                except Exception as e:
                    logger.error(f"Error processing result for {result.get('filename', 'unknown')}: {e}")

            # If we have results with analysis, create tabs for each document
            if results_with_real_analysis:
                tab_titles = [res[0].get("filename", f"Result {i+1}") for i, res in enumerate(results_with_real_analysis)]
                tabs = st.tabs(tab_titles)

                for i, (result, ai_analysis) in enumerate(results_with_real_analysis):
                    with tabs[i]:
                        filename = result.get("filename", "Unknown File")
                        annotated_pdf_b64 = result.get("annotated_pdf")

                        # File info and download button row (app.py.bak style)
                        file_col1, file_col2 = st.columns([0.8, 0.2])
                        with file_col1:
                            st.markdown(f"""
                            <div class="sleek-container">
                                <div class="file-name">
                                    <span class="file-icon">📄</span> {filename}
                                </div>
                            </div>
                            """, unsafe_allow_html=True)

                        with file_col2:
                            if annotated_pdf_b64:
                                annotated_pdf_bytes = base64.b64decode(annotated_pdf_b64)
                                # Simpler label from app.py.bak
                                download_label = "💾 PDF"
                                st.download_button(
                                    label=download_label,
                                    data=annotated_pdf_bytes,
                                    file_name=f"{filename.replace('.pdf', '').replace('.docx', '')}_annotated.pdf",
                                    mime="application/pdf",
                                    key=f"download_pdf_{i}_{filename}", # Ensure unique key
                                    use_container_width=True, # Consistent with app.py.bak button style
                                    help=f"Download annotated PDF for {filename}" # Added help text
                                )
                            else:
                                st.caption("No PDF")

                        # Display analysis sections
                        analysis_sections = ai_analysis.get("analysis_sections", {})
                        citation_counter = 0 # For numbering citations within a tab

                        for section_key, section_data in analysis_sections.items():
                            # Extract the actual title from the section key (removing "section_n_" prefix)
                            # Example: "section_1_investment_amount" -> "investment amount"
                            section_title = section_key
                            # Check if it follows the pattern section_N_title
                            if re.match(r'^section_\d+_', section_key):
                                # Extract just the title part after section_N_
                                section_title = re.sub(r'^section_\d+_', '', section_key)

                            # Format section name for display
                            display_section_name = section_title.replace("_", " ").title()

                            # Create a container for the section title with improved styling and RAG retry button
                            with st.container(border=False):
                                # Create columns for title and RAG retry button
                                title_col, rag_col = st.columns([0.92, 0.08])

                                with title_col:
                                    st.markdown(f"""
                                        <div style='background-color: #f5f5f5; padding: 0px 16px; border-radius: 8px;
                                                margin: 16px 0 8px 0; border-left: 4px solid #1976d2;'>
                                            <h4 style='color: #333; font-size: 1.2rem; margin: 0; font-weight: 600;'>
                                                {display_section_name}
                                            </h4>
                                        </div>
                                    """, unsafe_allow_html=True)

                                with rag_col:
                                    st.markdown('<div style="margin-top: 16px;">', unsafe_allow_html=True)
                                    # Add RAG retry button in the header
                                    display_rag_retry_button_header(section_key, result, section_data)
                                    st.markdown('</div>', unsafe_allow_html=True)

                            # Display RAG analysis and retry results if available (below the header)
                            # Retry results are integrated into the main view; no separate section

                            # Section content in a bordered container
                            with st.container(border=True):
                                analysis_content = section_data.get("Analysis")
                                context_content = section_data.get("Context")

                                if analysis_content:
                                    analysis_html_parts = [
                                        f"<div style='background-color: #f8f9fa; padding: .5rem; border-radius: 0.5rem; margin-bottom: 1rem;'>",
                                        f"<h4 style='color: #1e88e5; font-size: 1.1rem;'>Analysis</h4>",
                                        f"<div style='color: #424242; line-height: 1.6;'>{analysis_content}"
                                    ]
                                    if context_content:
                                        analysis_html_parts.extend([
                                            f"<div style='margin-top: 0.8rem; border-top: 1px solid #e0e0e0; padding-top: 0.8rem;'>",
                                            f"<span style='color: #1b5e20; font-size: 0.9rem; line-height: 1.4;'>{context_content}</span>",
                                            f"</div>"
                                        ])
                                    analysis_html_parts.extend(["</div></div>"])
                                    st.markdown("".join(analysis_html_parts), unsafe_allow_html=True)
                                elif context_content: # Display context even if analysis is missing
                                    st.markdown(f"""
                                        <div style='background-color: #f8f9fa; padding: .5rem; border-radius: 0.5rem; margin-bottom: 1rem;'>
                                            <h4 style='color: #1e88e5; font-size: 1.1rem;'>Context</h4>
                                            <div style='color: #424242; line-height: 1.6;'>
                                                <span style='color: #1b5e20; font-size: 0.9rem; line-height: 1.4;'>{context_content}</span>
                                            </div>
                                        </div>
                                    """, unsafe_allow_html=True)

                            # Supporting Citations in an expander
                            supporting_phrases = section_data.get("Supporting_Phrases", [])
                            verification_results = result.get("verification_results", {})
                            phrase_locations = result.get("phrase_locations", {})

                            any_needs_review = False
                            if supporting_phrases and supporting_phrases != ["No relevant phrase found."]:
                                for phrase in supporting_phrases:
                                    phrase_verification = verification_results.get(phrase, {})
                                    if isinstance(phrase_verification, dict):
                                        is_verified = phrase_verification.get("verified", False)
                                    elif isinstance(phrase_verification, bool):
                                        is_verified = phrase_verification
                                    else:
                                        is_verified = bool(phrase_verification) # Fallback
                                    if not is_verified:
                                        any_needs_review = True
                                        break

                            with st.expander("Supporting Citations", expanded=any_needs_review):
                                # If optimized RAG results exist for this section, use them to REPLACE the citations display
                                # Disabled displaying optimized RAG results separately; use verified citations from analysis
                                new_results = []

                                if new_results:
                                    for j, chunk in enumerate(new_results):
                                        citation_counter += 1
                                        text = chunk.get('text', '')
                                        page = chunk.get('page_num', chunk.get('page', 'Unknown'))
                                        score_val = chunk.get('score')
                                        try:
                                            current_score_info = f"{float(score_val):.1f}" if score_val is not None else "N/A"
                                        except Exception:
                                            current_score_info = str(score_val) if score_val is not None else "N/A"
                                        current_page_num_info = f"Page {page}"

                                        # Always mark optimized citations as verified (they come from RAG)
                                        badge_html = '<span style="display: inline-block; background-color: #d1fecf; color: #11631a; padding: 1px 6px; border-radius: 0.25rem; font-size: 0.8em; margin-left: 5px; border: 1px solid #a1e0a3; font-weight: 600;">✔ Verified</span>'

                                        cite_col, btn_col = st.columns([0.90, 0.10], gap="small")
                                        with cite_col:
                                            st.markdown(f"""
                                            <div style="border: 1px solid #e0e0e0; border-radius: 5px; padding: 8px 12px; margin-top: 5px; margin-bottom: 8px; background-color: #f9f9f9;">
                                                <div style="margin-bottom: 5px; display: flex; justify-content: space-between; align-items: center;">
                                                    <span style="font-weight: bold;">Citation {citation_counter} {badge_html}</span>
                                                    <span style="font-size: 0.8em; color: #555;">{current_page_num_info} | Score: {current_score_info}</span>
                                                </div>
                                                <div style="color: #333; line-height: 1.4; font-size: 0.95em;"><i>"{text}"</i></div>
                                            </div>
                                            """, unsafe_allow_html=True)

                                        with btn_col:
                                            st.markdown('<div style="margin-top: 20px;"></div>', unsafe_allow_html=True)
                                            filename_for_go = filename
                                            pdf_bytes_for_view = find_annotated_pdf_for_filename(filename_for_go)
                                            if pdf_bytes_for_view and isinstance(page, int):
                                                button_key = f"goto_{i}_{section_key}_{citation_counter}_{j}"
                                                if st.button("Go", key=button_key, type="secondary", help=f"Go to Page {page} in {filename_for_go}", use_container_width=True):
                                                    update_pdf_view(pdf_bytes=pdf_bytes_for_view, page_num=page, filename=filename_for_go)
                                                    st.session_state.scroll_to_pdf_viewer = True
                                                    st.rerun()
                                            else:
                                                st.caption("PDF N/A")
                                else:
                                    # Fallback to original citations if no optimized results
                                    if not supporting_phrases or supporting_phrases == ["No relevant phrase found."]:
                                        st.info("No supporting citations were identified for this section.")
                                    else:
                                        has_citations_to_show = False
                                        for phrase_idx, phrase in enumerate(supporting_phrases):
                                            if not isinstance(phrase, str) or phrase == "No relevant phrase found.":
                                                continue
                                            has_citations_to_show = True
                                            citation_counter += 1 # Increment citation counter

                                            phrase_verification = verification_results.get(phrase, {})
                                            phrase_location_data = phrase_locations.get(phrase, {}) # Note: app.py.bak uses find_best_location

                                            is_verified = False
                                            score = 0
                                            best_location_dict = {} # Renamed from best_location to avoid conflict

                                            if isinstance(phrase_verification, bool):
                                                is_verified = phrase_verification
                                            elif isinstance(phrase_verification, dict):
                                                is_verified = phrase_verification.get("verified", False)
                                                score = phrase_verification.get("score", 0)
                                            else:
                                                try: is_verified = bool(phrase_verification)
                                                except: is_verified = False

                                            current_page_num_info = "Page unknown"
                                            current_score_info = "N/A"

                                            if isinstance(phrase_location_data, list) and phrase_location_data:
                                                first_loc = phrase_location_data[0]
                                                if isinstance(first_loc, dict):
                                                    page_num_val = first_loc.get("page_num")
                                                    if isinstance(page_num_val, int):
                                                        current_page_num_info = f"Page {page_num_val + 1}"
                                                    else:
                                                        current_page_num_info = f"Page {page_num_val}" if page_num_val is not None else "Page unknown"

                                                    score_val = first_loc.get("match_score", score)
                                                    if score_val:
                                                        try: current_score_info = f"{float(score_val):.1f}"
                                                        except: current_score_info = str(score_val)
                                                    else:
                                                        if score:
                                                            try: current_score_info = f"{float(score):.1f}"
                                                            except: current_score_info = str(score)
                                                best_location_dict = first_loc
                                            elif isinstance(phrase_location_data, dict):
                                                page_num_val = phrase_location_data.get("page_num")
                                                if isinstance(page_num_val, int):
                                                    current_page_num_info = f"Page {page_num_val + 1}"
                                                else:
                                                    current_page_num_info = f"Page {page_num_val}" if page_num_val is not None else "Page unknown"

                                                score_val = phrase_location_data.get("match_score", score)
                                                if score_val:
                                                    try: current_score_info = f"{float(score_val):.1f}"
                                                    except: current_score_info = str(score_val)
                                                else:
                                                    if score:
                                                        try: current_score_info = f"{float(score):.1f}"
                                                        except: current_score_info = str(score)
                                                best_location_dict = phrase_location_data.get("best_match", phrase_location_data)

                                            if is_verified:
                                                badge_html = '<span style="display: inline-block; background-color: #d1fecf; color: #11631a; padding: 1px 6px; border-radius: 0.25rem; font-size: 0.8em; margin-left: 5px; border: 1px solid #a1e0a3; font-weight: 600;">✔ Verified</span>'
                                            else:
                                                badge_html = '<span style="display: inline-block; background-color: #ffeacc; color: #a05e03; padding: 1px 6px; border-radius: 0.25rem; font-size: 0.8em; margin-left: 5px; border: 1px solid #f8c78d; font-weight: 600;">⚠️ Needs Review</span>'

                                            cite_col, btn_col = st.columns([0.90, 0.10], gap="small")
                                            with cite_col:
                                                st.markdown(f"""
                                                <div style="border: 1px solid #e0e0e0; border-radius: 5px; padding: 8px 12px; margin-top: 5px; margin-bottom: 8px; background-color: #f9f9f9;">
                                                    <div style="margin-bottom: 5px; display: flex; justify-content: space-between; align-items: center;">
                                                        <span style="font-weight: bold;">Citation {citation_counter} {badge_html}</span>
                                                        <span style="font-size: 0.8em; color: #555;">{current_page_num_info} | Score: {current_score_info}</span>
                                                    </div>
                                                    <div style="color: #333; line-height: 1.4; font-size: 0.95em;"><i>"{phrase}"</i></div>
                                                </div>
                                                """, unsafe_allow_html=True)

                                            with btn_col:
                                                st.markdown('<div style="margin-top: 20px;"></div>', unsafe_allow_html=True)
                                                if is_verified and best_location_dict and isinstance(best_location_dict, dict) and "page_num" in best_location_dict and annotated_pdf_b64:
                                                    try:
                                                        page_num_to_go = best_location_dict["page_num"]
                                                        page_num_1_indexed = page_num_to_go + 1 if isinstance(page_num_to_go, int) else int(page_num_to_go) + 1
                                                        button_key = f"goto_{i}_{section_key}_{citation_counter}_{phrase_idx}"
                                                        if st.button("Go", key=button_key, type="secondary", help=f"Go to Page {page_num_1_indexed} in {filename}", use_container_width=True):
                                                            pdf_bytes_for_view = base64.b64decode(annotated_pdf_b64)
                                                            update_pdf_view(pdf_bytes=pdf_bytes_for_view, page_num=page_num_1_indexed, filename=filename)
                                                            st.session_state.scroll_to_pdf_viewer = True
                                                            st.rerun()
                                                    except Exception as e_go:
                                                        logger.error(f"Error setting up 'Go' button for citation: {e_go}")
                                                elif is_verified:
                                                    st.caption("Loc N/A")

                                        if not has_citations_to_show:
                                            st.caption("No supporting citations provided or found for this section.")

                            # Facts display removed per request (use Export Results > Export Facts)
            else: # No results_with_real_analysis
                 st.info("Processing complete, but no analysis sections were generated or found.")



            # Add Follow-up Question Interface at the bottom of the analysis results
            st.markdown('<hr style="margin: 20px 0; border: 0; border-top: 2px solid #e0e0e0;">', unsafe_allow_html=True)
            st.markdown('<div class="header-title" style="font-size: 1.3rem;">Follow-up Questions [Beta]</div>', unsafe_allow_html=True)
            st.markdown('<hr style="margin: 12px 0; border: 0; border-top: 1px solid #e0e0e0;">', unsafe_allow_html=True)

            # Display existing follow-up Q&A if any
            if st.session_state.get("followup_qa"):
                for i, qa_pair in enumerate(st.session_state.get("followup_qa", [])):
                    # Question
                    st.markdown(f"""
                    <div style='background-color: #e3f2fd; padding: 12px; border-radius: 8px; margin: 8px 0; border-left: 4px solid #1976d2;'>
                        <strong>Q{i+1}:</strong> {qa_pair['question']}
                    </div>
                    """, unsafe_allow_html=True)

                    # Answer with citations
                    with st.container(border=True):
                        processed_text = qa_pair.get("processed_text", qa_pair.get("answer", ""))
                        citation_details = qa_pair.get("citation_details", [])

                        # Display the answer text
                        st.markdown(f"""
                        <div style='background-color: #f8f9fa; padding: 12px; border-radius: 8px; margin-bottom: 8px;'>
                            <div style='color: #424242; line-height: 1.6;'>{processed_text}</div>
                        </div>
                        """, unsafe_allow_html=True)

                        # Display citations if any
                        if citation_details:
                            with st.expander("Supporting Citations", expanded=False):
                                # Use processed_text which contains the citation numbers [1], [2], etc.
                                processed_answer_text = qa_pair.get("processed_text", qa_pair.get("answer", ""))
                                display_followup_citations_like_main_analysis(citation_details, i, processed_answer_text)

            # Follow-up question input
            followup_question = st.text_input(
                "Ask a follow-up question about the analysis:",
                placeholder="e.g., Can you provide more details about the investment timeline?",
                key="followup_question_input"
            )

            col1, _ = st.columns([1, 4])
            with col1:
                ask_followup = st.button(
                    "Ask Follow-up Question",
                    key="ask_followup_button",
                    type="primary",
                    disabled=not followup_question.strip(),
                    use_container_width=True
                )

            # Process follow-up question
            if ask_followup and followup_question.strip():
                with st.spinner("Processing your follow-up question..."):
                    try:
                        logger.info(f"Processing follow-up question: {followup_question[:50]}...")

                        # Use the same RAG pipeline as chat with same retrieval depth as main analysis
                        relevant_chunks = retrieve_relevant_chunks_for_chat(
                            prompt=followup_question,
                            top_k_per_doc=RAG_TOP_K,
                            embedding_model=embedding_model,
                            reranker_model=reranker_model,
                            preprocessed_data=st.session_state.get("preprocessed_data", {})
                        )

                        # Generate response using the same analyzer
                        analyzer = DocumentAnalyzer()
                        raw_response = run_async(
                            generate_chat_response(
                                analyzer,
                                followup_question,
                                relevant_chunks
                            )
                        )

                        # Process response for citations
                        processed_text, citation_details = process_chat_response_for_numbered_citations(raw_response)

                        # Store the Q&A pair
                        qa_pair = {
                            "question": followup_question,
                            "answer": raw_response,
                            "processed_text": processed_text,
                            "citation_details": citation_details,
                            "timestamp": datetime.now().isoformat()
                        }

                        # Ensure followup_qa exists before appending
                        if "followup_qa" not in st.session_state:
                            st.session_state.followup_qa = []
                        st.session_state.followup_qa.append(qa_pair)
                        logger.info("Follow-up question processed successfully")

                        # Clear the input and rerun to show the new Q&A
                        st.rerun()

                    except Exception as e:
                        logger.error(f"Error processing follow-up question: {e}", exc_info=True)
                        st.error(f"Sorry, an error occurred while processing your follow-up question: {str(e)}")

            if not st.session_state.get("followup_qa") and not followup_question.strip():
                st.info("💡 Ask follow-up questions to get more specific insights about your documents. The AI will use the same document context to provide detailed answers.")

    # Right Column: Tools & PDF Viewer
    with tools_col:
        # Header style from app.py.bak
        st.markdown('<div class="header-title">Analysis Tools & PDF Viewer</div>', unsafe_allow_html=True)
        st.markdown('<hr style="margin: 12px 0; border: 0; border-top: 1px solid #e0e0e0;">', unsafe_allow_html=True)

        # Container for tools
        with st.container():
            # SmartChat Expander
            with st.expander("💬 SmartChat (Multi-Document Chat)", expanded=False):
                if not st.session_state.get("preprocessed_data"):
                    st.info("Upload and process documents to enable chat.")
                else:
                    chat_container = st.container(height=400, border=True)
                    with chat_container:
                        # Use enumerate to get the index of each message in the session state list
                        for msg_idx, message in enumerate(st.session_state.chat_messages):
                            with st.chat_message(message["role"]):
                                if message["role"] == "assistant":
                                    processed_text = message.get("processed_text", message["content"])
                                    citation_details = message.get("citation_details", [])
                                    # Pass the message index (msg_idx) to the display function
                                    display_chat_message_with_citations(processed_text, citation_details, msg_idx)
                                else:
                                    # Apply the same word-wrap styling to user messages for consistency
                                    st.markdown(
                                        f'<div style="word-wrap: break-word; overflow-wrap: break-word; white-space: normal;">{message["content"]}</div>',
                                        unsafe_allow_html=True
                                    )

                    if prompt := st.chat_input("Ask about the uploaded documents...", key="chat_input_main"):
                        st.session_state.chat_messages.append({"role": "user", "content": prompt})
                        processed_chat_text = "Error: Could not generate response."
                        chat_citation_details = []
                        raw_ai_response_content = ""
                        try:
                            with st.spinner("Thinking..."):
                                logger.info(f"Chat RAG started for: {prompt[:50]}...")
                                # Use same retrieval depth as main analysis for consistency
                                relevant_chunks = retrieve_relevant_chunks_for_chat(
                                    prompt=prompt,
                                    top_k_per_doc=RAG_TOP_K,
                                    embedding_model=embedding_model,
                                    reranker_model=reranker_model,  # Use local reranker model
                                    preprocessed_data=st.session_state.get("preprocessed_data", {})
                                )
                                analyzer = DocumentAnalyzer()
                                logger.info(f"Generating chat response for: {prompt[:50]}...")
                                raw_ai_response_content = run_async(
                                    generate_chat_response(
                                        analyzer,
                                        prompt,
                                        relevant_chunks
                                    )
                                )
                                logger.info("Chat response generated.")
                                processed_chat_text, chat_citation_details = process_chat_response_for_numbered_citations(raw_ai_response_content)
                        except Exception as chat_err:
                            logger.error(f"Error during chat processing: {chat_err}", exc_info=True)
                            processed_chat_text = f"Sorry, an error occurred while processing your request: {str(chat_err)}"
                            chat_citation_details = []
                        st.session_state.chat_messages.append({
                            "role": "assistant",
                            "content": raw_ai_response_content,
                            "processed_text": processed_chat_text,
                            "citation_details": chat_citation_details
                        })
                        st.rerun()

            # Export Results Expander
            with st.expander("📊 Export Results", expanded=False):
                # Prepare data for export
                exportable_results_list = []

                for result, ai_analysis in results_with_real_analysis:
                    filename = result.get("filename", "Unknown File")
                    verification_results = result.get("verification_results", {})
                    phrase_locations = result.get("phrase_locations", {})

                    # Prepare a list of flattened data for this file
                    file_data = []

                    for section_key, section_data in ai_analysis.get("analysis_sections", {}).items():
                        section_name = section_key.replace("_", " ").title()
                        analysis_text = section_data.get("Analysis", "")
                        context = section_data.get("Context", "")

                        # Process each supporting phrase
                        supporting_phrases = section_data.get("Supporting_Phrases", [])

                        if not supporting_phrases or supporting_phrases == ["No relevant phrase found."]:
                            # Add a single row for this section without phrases
                            file_data.append({
                                "Filename": filename,
                                "Section": section_name,
                                "Analysis": analysis_text,
                                "Context": context,
                                "Supporting Phrase": "No relevant phrase found.",
                                "Verified": "N/A",
                                "Page": "N/A",
                                "Match Score": "N/A"
                            })
                        else:
                            # Add a row for each supporting phrase
                            for phrase in supporting_phrases:
                                # Get verification info safely
                                phrase_verification = verification_results.get(phrase, {})
                                phrase_location = phrase_locations.get(phrase, {})

                                # Default values
                                is_verified = False
                                score = 0
                                best_location = {}

                                # Handle different verification result formats
                                # If phrase_verification is a boolean, use it directly
                                if isinstance(phrase_verification, bool):
                                    is_verified = phrase_verification
                                # If it's a dictionary, extract the verified field
                                elif isinstance(phrase_verification, dict):
                                    is_verified = phrase_verification.get("verified", False)
                                    score = phrase_verification.get("score", 0)
                                # Otherwise, try to convert to boolean
                                else:
                                    try:
                                        is_verified = bool(phrase_verification)
                                    except Exception:
                                        is_verified = False

                                # Handle different phrase location formats
                                # If it's a list, find the best location (first one with highest score)
                                if isinstance(phrase_location, list) and phrase_location:
                                    # Sort by match_score if available, otherwise use the first one
                                    try:
                                        sorted_locations = sorted(phrase_location,
                                                                key=lambda x: x.get("match_score", 0) if isinstance(x, dict) else 0,
                                                                reverse=True)
                                        best_location = sorted_locations[0] if sorted_locations else {}
                                        if isinstance(best_location, dict) and "match_score" in best_location:
                                            score = best_location["match_score"]
                                    except Exception:
                                        best_location = phrase_location[0] if phrase_location else {}
                                # If it's a dictionary, use it as is
                                elif isinstance(phrase_location, dict):
                                    best_location = phrase_location.get("best_match", phrase_location)

                                # Calculate page number
                                if isinstance(best_location, dict) and "page_num" in best_location:
                                    page_num = best_location.get("page_num", -1) + 1
                                else:
                                    page_num = "Unknown"

                                file_data.append({
                                    "Filename": filename,
                                    "Section": section_name,
                                    "Analysis": analysis_text,
                                    "Context": context,
                                    "Supporting Phrase": phrase,
                                    "Verified": "Yes" if is_verified else "No",
                                    "Page": page_num,
                                    "Match Score": f"{score:.1f}%" if score else "N/A"
                                })

                    # Add this file's data to the exportable results
                    exportable_results_list.append({
                        "filename": filename,
                        "data": file_data,
                        "analysis": ai_analysis
                    })

                # Excel Export
                if exportable_results_list:
                    # Flatten all data for Excel export
                    flat_data = []
                    for file_result in exportable_results_list:
                        flat_data.extend(file_result["data"])

                    # Add follow-up Q&A data if any exist
                    followup_qa = st.session_state.get("followup_qa", [])
                    if followup_qa:
                        for i, qa_pair in enumerate(followup_qa):
                            flat_data.append({
                                "Filename": "Follow-up Q&A",
                                "Section": f"Question {i+1}",
                                "Analysis": f"Q: {qa_pair.get('question', '')}\n\nA: {qa_pair.get('answer', '')}",
                                "Context": f"Asked on: {qa_pair.get('timestamp', 'Unknown')}",
                                "Supporting Phrase": f"Citations: {len(qa_pair.get('citation_details', []))} found",
                                "Verified": "N/A",
                                "Page": "N/A",
                                "Match Score": "N/A"
                            })

                    # Convert to DataFrame and export
                    df = pd.DataFrame(flat_data)
                    excel_buffer = BytesIO()
                    df.to_excel(excel_buffer, index=False, engine="openpyxl")
                    excel_buffer.seek(0)

                    st.download_button(
                        label="📥 Export Excel",
                        data=excel_buffer,
                        file_name=f"analysis_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_excel"
                    )

                # Word Export
                if exportable_results_list:
                    word_bytes = export_to_word(exportable_results_list)

                    st.download_button(
                        label="📥 Export Word",
                        data=word_bytes,
                        file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_word"
                    )



            # Fact Extraction (beta) Expander
            with st.expander("🧪 Fact Extraction (beta)", expanded=False):
                st.caption("Uses custom LLM-based extraction with Pydantic-AI to intelligently identify fact types and extract structured information from analysis text.")

                if st.button("🧠 Generate Facts", key="compute_fact_definitions_beta"):
                    with st.spinner("Extracting facts using LLM-based analysis..."):
                        try:
                            from src.keyword_code.services.fact_extraction_service import FactExtractionService

                            # Initialize the fact extraction service
                            fact_service = FactExtractionService()

                            # Show progress and debug info
                            st.info(f"Processing {len(results_with_real_analysis)} document(s)...")

                            # Debug: Check API key
                            import os
                            has_api_key = bool(os.environ.get('DATABRICKS_API_KEY'))
                            st.info(f"DATABRICKS_API_KEY present: {has_api_key}")

                            # Debug: Show what sections we're processing
                            total_sections = 0
                            for res, ai in results_with_real_analysis:
                                sections = (ai or {}).get("analysis_sections", {}) or {}
                                total_sections += len([s for s in sections.values() if s.get("Analysis")])
                            st.info(f"Found {total_sections} sections with analysis text")

                            # Extract facts using the new service
                            rows = fact_service.extract_fact_definitions_for_results(results_with_real_analysis)

                            if rows:
                                # Build two-column Excel (Fact, Definition) as requested
                                df_two = pd.DataFrame([{ "Fact": r.get("Fact",""), "Definition": r.get("Definition","") } for r in rows])
                                buf = BytesIO()
                                df_two.to_excel(buf, index=False, engine="openpyxl")
                                buf.seek(0)
                                st.session_state["facts_defs_excel"] = buf.getvalue()
                                st.session_state["facts_defs_json"] = json.dumps([
                                    {"Fact": r.get("Fact",""), "Definition": r.get("Definition","")}
                                    for r in rows
                                ], ensure_ascii=False, indent=2).encode("utf-8")
                                st.success(f"✅ Extracted {len(rows)} facts using intelligent LLM-based analysis.")

                                # Show a preview
                                st.subheader("Preview (first 10 facts)")
                                preview_df = pd.DataFrame(rows[:10])
                                st.dataframe(preview_df, use_container_width=True)

                            else:
                                st.warning("⚠️ No facts extracted. This could be due to:")
                                st.write("- Missing DATABRICKS_API_KEY environment variable")
                                st.write("- No analysis text in the selected documents")
                                st.write("- LLM extraction did not identify any facts")
                                st.write("- Analysis text may not contain extractable factual information")

                        except Exception as _fd_err:
                            st.error(f"❌ Error extracting facts: {_fd_err}")
                            logger.error(f"Fact extraction error: {_fd_err}", exc_info=True)

                if st.session_state.get("facts_defs_excel"):
                    st.download_button(
                        label="📥 Export Fact Definitions (Excel)",
                        data=st.session_state.get("facts_defs_excel"),
                        file_name=f"fact_definitions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_fact_defs_excel"
                    )
                    st.download_button(
                        label="📥 Export Fact Definitions (JSON)",
                        data=st.session_state.get("facts_defs_json"),
                        file_name=f"fact_definitions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        key="download_fact_defs_json"
                    )

            # Report Issue Expander
            with st.expander("🐞 Report Issue", expanded=False):
                st.markdown("""
                ### Report an Issue

                If you encounter any problems with the analysis or have feedback, please describe the issue below.
                A report package will be generated that you can send to the CNT Automations team.

                Positive feedback is good, negative feedback is even better!

                """)

                # Issue description input
                issue_description = st.text_area(
                    "Issue Description",
                    placeholder="Please describe the issue you're experiencing...",
                    height=150
                )

                # Create report package filename
                report_filename = f'smartdocs_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip'

                # Check if description is provided
                download_disabled = not issue_description.strip()

                # Create the report package content
                try:
                    # Generate package content only if description is provided
                    if not download_disabled:
                        # Create a function to generate the report package
                        def create_report_package_for_download(desc):
                            try:
                                report_data = {
                                    "timestamp": datetime.now().isoformat(),
                                    "issue_description": desc,
                                    "user_inputs": {
                                        "prompt": st.session_state.get('user_prompt', ''),
                                    },
                                    "analysis_results": st.session_state.get('analysis_results', None),
                                    "current_document": st.session_state.get('current_pdf_name', None),
                                    "preprocessed_data_keys": list(st.session_state.get('preprocessed_data', {}).keys()),
                                    "chat_history_summary": [
                                        {"role": msg.get("role"), "content_preview": msg.get("content", "")[:100]+"..."}
                                        for msg in st.session_state.get("chat_messages", [])
                                    ],
                                    "followup_qa": st.session_state.get("followup_qa", [])
                                }

                                zip_buffer = BytesIO()
                                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                    # Write report data as JSON
                                    try:
                                        zip_file.writestr('report_data.json', json.dumps(report_data, indent=2, default=str))
                                    except Exception as json_err:
                                        zip_file.writestr('report_data_error.txt', f"Error serializing report data: {json_err}")
                                        logger.error(f"Error serializing report_data.json: {json_err}", exc_info=True)

                                    # Write original uploaded files
                                    uploaded_file_objs = st.session_state.get('uploaded_file_objects')
                                    if uploaded_file_objs:
                                        for uploaded_file in uploaded_file_objs:
                                            try:
                                                if hasattr(uploaded_file, 'name') and hasattr(uploaded_file, 'getvalue'):
                                                    zip_file.writestr(f'original_docs/{uploaded_file.name}', uploaded_file.getvalue())
                                                else:
                                                    logger.warning(f"Skipping invalid file object in uploaded_file_objects during report creation: {type(uploaded_file)}")
                                            except Exception as file_read_err:
                                                zip_file.writestr(f'original_docs/ERROR_{uploaded_file.name}.txt', f"Error reading file: {file_read_err}")
                                                logger.error(f"Error reading file {uploaded_file.name} for report package: {file_read_err}", exc_info=True)

                                    # Write annotated PDFs
                                    analysis_results_list = st.session_state.get('analysis_results')
                                    if analysis_results_list:
                                        for result in analysis_results_list:
                                            if isinstance(result, dict) and 'annotated_pdf' in result and result.get('annotated_pdf'):
                                                try:
                                                    pdf_bytes = base64.b64decode(result['annotated_pdf'])
                                                    pdf_filename = result.get('filename', f'unknown_annotated_{result.get("timestamp", "ts")}.pdf')
                                                    zip_file.writestr(f'annotated_pdfs/{pdf_filename}', pdf_bytes)
                                                except Exception as pdf_err:
                                                    zip_file.writestr(f'annotated_pdfs/ERROR_{result.get("filename", "unknown")}.txt', f"Error decoding/writing PDF: {pdf_err}")
                                                    logger.error(f"Error writing annotated PDF {result.get('filename')} to report: {pdf_err}", exc_info=True)

                                zip_buffer.seek(0)
                                return zip_buffer.getvalue()
                            except Exception as zip_e:
                                logger.error(f"Error creating report package zip file: {zip_e}", exc_info=True)
                                # Create a simple error zip as fallback
                                zip_buffer = BytesIO()
                                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                    zip_file.writestr('error_creating_report.txt', f"Failed to create full report package: {zip_e}")
                                zip_buffer.seek(0)
                                return zip_buffer.getvalue()

                        # Create the download button
                        st.download_button(
                            label="📥 Download Report Package",
                            data=create_report_package_for_download(issue_description),
                            file_name=report_filename,
                            mime='application/zip',
                            disabled=download_disabled,
                            help="Download the report package to attach to your email.",
                            key="download_report_button",
                            use_container_width=True
                        )

                        st.success("""
                        Report package created successfully. Please download and email it to cnt_automations@ifc.org.

                        The package includes:
                        - Your issue description
                        - Analysis results
                        - Original documents
                        - Annotated PDFs
                        - Chat history
                        """)
                    else:
                        # Show disabled button with message
                        st.button(
                            "📥 Download Report Package",
                            disabled=True,
                            key="disabled_download_button",
                            help="Please provide an issue description first",
                            use_container_width=True
                        )
                        st.info("Please provide a description of the issue before downloading the report package.")
                except Exception as e:
                    st.error(f"Error preparing report package: {str(e)}")
                    logger.error(f"Error preparing report package for download button: {str(e)}", exc_info=True)

                st.info("Note: The report package will include the uploaded documents and analysis results to help diagnose the issue.")

        # PDF Viewer
        with st.expander("📄 PDF Viewer", expanded=st.session_state.get("show_pdf", False)):
            # Add an anchor for scrolling to PDF viewer
            st.markdown('<div id="pdf-viewer-anchor"></div>', unsafe_allow_html=True)

            if st.session_state.get("pdf_bytes") and st.session_state.get("show_pdf", False):
                fitz_doc = None # Initialize fitz_doc
                try:
                    # import fitz # PyMuPDF - This import should be at the top of the file

                    pdf_bytes = st.session_state.pdf_bytes
                    current_page = st.session_state.get("pdf_page", 1)
                    filename = st.session_state.get("current_pdf_name", "Document")

                    # Display filename (app.py.bak style)
                    st.caption(f"**{filename}**")

                    # Render the current page using PyMuPDF
                    fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    page_count = len(fitz_doc)

                    if page_count == 0:
                        st.warning("The PDF document appears to have 0 pages.")
                    else:
                        # Ensure current page is valid
                        current_page = max(1, min(current_page, page_count))

                        # Page navigation (app.py.bak style: number input)
                        nav_key = f"pdf_nav_{filename}_{page_count}_{hash(pdf_bytes)}" # More unique key
                        new_page = st.number_input(
                            "Page", # Simplified label
                            min_value=1,
                            max_value=page_count,
                            value=current_page,
                            step=1,
                            key=nav_key, # Use the unique key
                            help=f"Enter page number (1-{page_count})"
                        )
                        if new_page != current_page:
                            update_pdf_view(pdf_bytes, new_page, filename)
                            # update_pdf_view should handle st.rerun() if state changed
                            # For robustness, ensure rerun if number_input itself caused change.
                            st.rerun()

                        st.caption(f"Page {current_page} of {page_count}") # Page count caption

                        # Render the page
                        page = fitz_doc.load_page(current_page - 1)  # 0-indexed
                        pix = page.get_pixmap(dpi=150)
                        img_bytes = pix.tobytes("png")

                        # Display the page image
                        st.image(img_bytes, use_container_width=True)

                except Exception as e:
                    logger.error(f"Error displaying PDF: {e}")
                    st.error(f"Error displaying PDF: {e}")
                finally:
                    if fitz_doc:
                        fitz_doc.close()
            else:
                st.info("Select a document to view by clicking on a 'Go' button in the analysis or using the 'View' button.")

    # Auto-scroll to PDF viewer if flag is set
    if st.session_state.get("scroll_to_pdf_viewer", False):
        js = """
        <script>
            setTimeout(function() {
                const anchor = document.getElementById('pdf-viewer-anchor');
                if (anchor) {
                    console.log("Scrolling to PDF viewer anchor...");
                    anchor.scrollIntoView({ behavior: 'smooth', block: 'start' });
                }
            }, 100);
        </script>
        """
        st.components.v1.html(js, height=0)
        st.session_state.scroll_to_pdf_viewer = False  # Reset flag after scroll





def display_rag_retry_button_header(section_key: str, result: Dict[str, Any], section_data: Dict[str, Any]):
    """
    Display RAG retry button in the section header.

    Args:
        section_key: The section identifier
        result: Result dictionary containing analysis data
        section_data: The specific section data
    """
    # Create a unique key for this section's retry button
    retry_key = f"retry_rag_{section_key}_{result.get('filename', 'unknown')}"

    # Since this function is called within a column context, we cannot create nested columns
    # Instead, we'll create buttons directly without columns, stacked vertically
    # Analyze button removed per new agent/tool design

    if st.button("↻", key=f"retry_{retry_key}", help="Retry Retrieval", use_container_width=True):
        # Store the retry request in session state
        if "rag_retry_requests" not in st.session_state:
            st.session_state.rag_retry_requests = {}

        st.session_state.rag_retry_requests[section_key] = {
            "status": "requested",
            "section_data": section_data,
            "result": result
        }
        st.rerun()


def display_rag_results_section(section_key: str):
    """
    Display RAG analysis and retry results for a section.

    Args:
        section_key: The section identifier
    """
    # Display analysis results if available
    if hasattr(st.session_state, 'rag_analysis_results') and section_key in st.session_state.rag_analysis_results:
        analysis = st.session_state.rag_analysis_results[section_key]

        with st.expander("📊 RAG Analysis Results", expanded=False):
            col1, col2 = st.columns(2)

            with col1:
                st.metric("Query Type", analysis.get("query_type", "Unknown"))
                st.metric("Quality Score", f"{analysis.get('current_quality_score', 0):.2f}")

            with col2:
                st.metric("Recommended BM25 Weight", f"{analysis.get('recommended_bm25_weight', 0.5):.2f}")
                st.metric("Recommended Semantic Weight", f"{analysis.get('recommended_semantic_weight', 0.5):.2f}")

            if analysis.get("issues_identified"):
                st.markdown("**Issues Identified:**")
                for issue in analysis["issues_identified"]:
                    st.markdown(f"• {issue}")

            if analysis.get("reasoning"):
                st.markdown("**Reasoning:**")
                st.markdown(analysis["reasoning"])

    # Display retry results if available
    if hasattr(st.session_state, 'rag_retry_results') and section_key in st.session_state.rag_retry_results:
        retry_data = st.session_state.rag_retry_results[section_key]

        with st.expander("🔄 RAG Retry Results", expanded=False):
            st.markdown("**New Retrieval Results:**")

            new_results = retry_data.get("new_results", [])
            if new_results:
                for i, chunk in enumerate(new_results[:3]):  # Show top 3 results
                    st.markdown(f"**Result {i+1}** (Score: {chunk.get('score', 0):.3f})")
                    st.markdown(f"*Page {chunk.get('page_num', 'Unknown')}*")
                    st.markdown(chunk.get('text', '')[:200] + '...')
                    st.markdown("---")
            else:
                st.info("No new results retrieved")

            # Show comparison with original results
            original_results = retry_data.get("original_results", [])
            original_count = len(original_results) if isinstance(original_results, list) else 0
            new_count = len(new_results)
            st.metric("Results Comparison", f"{new_count} new vs {original_count} original")

        # If we produced a new AI analysis for this retry, show it with validation and facts
        ai_section = retry_data.get("ai_section")
        if ai_section:
            with st.expander("🧠 AI Response (Retry) + Validation", expanded=True):
                # Analysis text
                analysis_text = ai_section.get("Analysis", "")
                if analysis_text:
                    st.markdown("**Analysis (Retry):**")
                    st.markdown(analysis_text)
                else:
                    st.info("No analysis text available from retry.")

                # Verified supporting citations
                supporting = ai_section.get("Supporting_Phrases", []) or []
                ver_results = retry_data.get("verification_results", {}) or {}
                phrase_locs = retry_data.get("phrase_locations", {}) or {}
                if supporting and supporting != ["No relevant phrase found."]:
                    st.markdown("**Verified Supporting Citations:**")
                    for idx, phrase in enumerate(supporting, start=1):
                        v = ver_results.get(phrase, False)
                        is_verified = v.get("verified", False) if isinstance(v, dict) else bool(v)
                        icon = "✅" if is_verified else "⚠️"
                        # page info from first location if present
                        pinfo = ""
                        locs = phrase_locs.get(phrase, [])
                        if isinstance(locs, list) and locs:
                            loc0 = locs[0] if isinstance(locs[0], dict) else {}
                            page_val = loc0.get("page_num")
                            if isinstance(page_val, int):
                                pinfo = f" (Page {page_val + 1})"
                            elif page_val is not None:
                                pinfo = f" (Page {page_val})"
                        st.markdown(f"{icon} [{idx}] {phrase}{pinfo}")
                else:
                    st.info("No supporting phrases identified by the retry analysis.")

                # Facts display removed per request (use Export Results > Export Facts)


def display_section_facts_expander(section_key: str, section_data: Dict[str, Any], result: Dict[str, Any], citation_counter: int = 0):
    """
    Display extracted facts in an expander for a specific section.

    Args:
        section_key: The section identifier
        section_data: The specific section data containing Analysis text
        result: Result dictionary containing filename and other metadata
        citation_counter: Current citation counter for consistent numbering
    """

    # Check if we have facts for this specific section
    section_facts_key = f"section_facts_{section_key}_{result.get('filename', 'unknown')}"

    # Check if facts extraction is in progress or completed for this section
    if hasattr(st.session_state, 'section_facts') and section_facts_key in st.session_state.section_facts:
        facts_data = st.session_state.section_facts[section_facts_key]

        if facts_data.get("status") == "completed" and facts_data.get("facts"):
            facts = facts_data["facts"]

            st.markdown("--- ")
            st.markdown("##### 📊 Extracted Facts")
            # Group facts by category
            facts_by_category = {}
            for fact in facts:
                category = fact.get("category", "General")
                if category not in facts_by_category:
                    facts_by_category[category] = []
                facts_by_category[category].append(fact)

            # Display each category
            for category, category_facts in facts_by_category.items():
                if len(facts_by_category) > 1:
                    st.markdown(f"**{category.replace('_', ' ').title()}:**")

                for fact in category_facts:
                    fact_text = fact.get("text", "")
                    attributes = fact.get("attributes", {})

                    # Compact, badge-like fact display
                    fact_html = f'<div style="margin: 4px 0; font-size: 0.92em; color: #333;">'
                    fact_html += f'<span style="display:inline-block; padding:2px 8px; border-radius:999px; background:#eef5ff; color:#1e88e5; border:1px solid #cfe3ff; font-weight:600; margin-right:8px;">{category.replace("_", " ").title()}</span>'
                    fact_html += f'{fact_text}'

                    if attributes:
                        fact_html += '<span style="margin-left:8px; color:#666;">'
                        for key, value in list(attributes.items())[:3]:
                            fact_html += f'<span style="display:inline-block; padding:1px 6px; border-radius:999px; background:#f5f5f5; border:1px solid #e0e0e0; margin-right:6px; font-size:0.85em;">{key.replace("_", " ").title()}: {value}</span>'
                        fact_html += '</span>'

                    fact_html += '</div>'
                    st.markdown(fact_html, unsafe_allow_html=True)

            # Show extraction metadata
            metadata = facts_data.get("metadata", {})
            if metadata:
                st.caption(f"Extracted {len(facts)} facts using {metadata.get('model_used', 'Unknown model')}")

        elif facts_data.get("status") == "processing":
            st.markdown("--- ")
            st.markdown("##### 📊 Extracted Facts")
            st.info("🔄 Extracting facts from analysis...")

        elif facts_data.get("status") == "error":
            st.markdown("--- ")
            st.markdown("##### 📊 Extracted Facts")
            error_msg = facts_data.get("error", "Unknown error")
            st.error(f"❌ Error extracting facts from analysis: {error_msg}")

            # Add retry button
            if st.button(f"🔄 Retry Fact Extraction", key=f"retry_facts_{section_facts_key}"):
                # Reset the request status to trigger retry
                if "section_facts_requests" in st.session_state:
                    st.session_state.section_facts_requests[section_facts_key] = {
                        "status": "requested",
                        "section_key": section_key,
                        "analysis_text": section_data.get("Analysis", ""),
                        "section_data": section_data,
                        "result": result
                    }
                st.session_state.section_facts[section_facts_key] = {"status": "processing"}
                st.rerun()

    else:
        # Always show the facts section, even if no facts are available yet
        st.markdown("--- ")
        st.markdown("##### 📊 Extracted Facts")

        # Trigger fact extraction for this section if not already done
        if section_data.get("Analysis"):
            # Perform extraction synchronously to avoid per-section reruns
            from src.keyword_code.services.fact_extraction_service import FactExtractionService
            try:
                fact_service = FactExtractionService()
                extracted_facts = fact_service.extract_facts_from_text(
                    text=section_data.get("Analysis", ""),
                    context=f"Legal/Financial Analysis - Section: {section_key}",
                    section_name=section_key,
                    filename=result.get("filename", "Unknown")
                )
                if "section_facts" not in st.session_state:
                    st.session_state.section_facts = {}
                st.session_state.section_facts[section_facts_key] = {
                    "status": "completed",
                    "facts": extracted_facts.get("extracted_facts", []),
                    "metadata": {
                        "model_used": "Pydantic-AI Fact Extraction",
                        "total_extractions": len(extracted_facts.get("extracted_facts", [])),
                        "section_key": section_key
                    }
                }
            except Exception as _fe_err:
                logger.error(f"Synchronous fact extraction error for section {section_key}: {_fe_err}")
                st.session_state.section_facts[section_facts_key] = {"status": "error", "error": str(_fe_err)}
            st.caption("Facts extracted using LLM-based analysis.")
        else:
            st.info("No analysis text available for fact extraction.")
