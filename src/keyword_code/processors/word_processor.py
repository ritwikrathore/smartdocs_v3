"""
Word document processing functionality.
"""

from io import BytesIO
from typing import Optional
import fitz  # PyMuPDF
from docx import Document as DocxDocument  # Renamed to avoid conflict
from ..config import logger


class WordProcessor:
    """Handles Word document conversion to PDF."""

    def __init__(self, docx_bytes: bytes):
        if not isinstance(docx_bytes, bytes):
            raise ValueError("docx_bytes must be of type bytes")
        self.docx_bytes = docx_bytes
        logger.info(f"WordProcessor initialized with {len(docx_bytes)} bytes.")

    def convert_to_pdf_bytes(self) -> Optional[bytes]:
        """Converts the DOCX to PDF using only built-in PyMuPDF fonts, with robust line/page breaking."""
        try:
            doc = DocxDocument(BytesIO(self.docx_bytes))
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        lines.append(" | ".join(row_text))
            if not lines:
                logger.warning("No text extracted from DOCX file.")
                return self.create_minimal_empty_pdf()

            pdf_doc = fitz.open()
            margin = 50
            fontsize = 10
            line_height = fontsize * 1.5
            page_width, page_height = 595, 842  # A4 in points
            max_width = page_width - 2 * margin

            # Try fonts in order of reliability
            for fontname in ["cour", "helv", "times"]:
                try:
                    # Prepare lines for PDF (wrap long lines)
                    wrapped_lines = []
                    dummy_page = pdf_doc.new_page(width=page_width, height=page_height)
                    for line in lines:
                        while line:
                            for i in range(len(line), 0, -1):
                                w = fitz.get_text_length(line[:i], fontname=fontname, fontsize=fontsize)
                                if w <= max_width:
                                    break
                            wrapped_lines.append(line[:i])
                            line = line[i:]
                    pdf_doc.delete_page(0)  # Remove dummy page

                    # Write lines to PDF, paginating as needed
                    page = pdf_doc.new_page(width=page_width, height=page_height)
                    y = margin
                    for line in wrapped_lines:
                        if y + line_height > page_height - margin:
                            page = pdf_doc.new_page(width=page_width, height=page_height)
                            y = margin
                        page.insert_text((margin, y), line, fontname=fontname, fontsize=fontsize)
                        y += line_height

                    # Successfully created PDF with this font
                    logger.info(f"Successfully converted DOCX to PDF using font '{fontname}'")
                    return pdf_doc.tobytes()
                except Exception as font_err:
                    logger.warning(f"Failed to create PDF with font '{fontname}': {font_err}")
                    # Try next font
                    continue

            # If we get here, all fonts failed
            logger.error("All font attempts failed for DOCX to PDF conversion.")
            return self.create_minimal_empty_pdf()

        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {str(e)}", exc_info=True)
            return self.create_minimal_empty_pdf()

    def create_minimal_empty_pdf(self) -> bytes:
        """Creates a minimal PDF with an error message when conversion fails."""
        try:
            pdf_doc = fitz.open()
            page = pdf_doc.new_page(width=595, height=842)  # A4
            error_text = "Error: Could not convert document. The file may be corrupted or in an unsupported format."
            page.insert_text((50, 50), error_text, fontname="helv", fontsize=12, color=(1, 0, 0))
            return pdf_doc.tobytes()
        except Exception as e:
            logger.error(f"Failed to create minimal error PDF: {e}")
            # Create an absolutely minimal PDF as last resort
            return b"%PDF-1.7\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj 2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj 3 0 obj<</Type/Page/MediaBox[0 0 595 842]/Parent 2 0 R>>endobj\nxref\n0 4\n0000000000 65535 f\n0000000010 00000 n\n0000000053 00000 n\n0000000102 00000 n\ntrailer<</Size 4/Root 1 0 R>>\nstartxref\n169\n%%EOF"
