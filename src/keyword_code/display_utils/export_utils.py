"""
Export utilities for analysis results.
"""

import streamlit as st
import base64
import json
import zipfile
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from typing import Dict, List, Any, Tuple
import fitz
from ..config import logger
from .pdf_utils import restore_original_bytes_if_needed


def create_report_package_content(issue_description: str, results: List[Tuple[Dict[str, Any], Dict[str, Any]]]) -> bytes:
    """
    Create a ZIP package containing all relevant information for issue reporting.

    Args:
        issue_description: User's description of the issue
        results: List of tuples containing (result, ai_analysis) pairs

    Returns:
        bytes: The ZIP file as bytes
    """
    # Create a BytesIO object to store the ZIP file
    zip_buffer = BytesIO()

    # Create a ZIP file
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        # Add report metadata
        report_data = {
            "issue_description": issue_description,
            "timestamp": datetime.now().isoformat(),
            "user_inputs": {
                "prompt": st.session_state.get("user_prompt", ""),
                "uploaded_files": list(st.session_state.get("preprocessed_data", {}).keys())
            },
            "chat_history": st.session_state.get("chat_messages", []),
            "followup_qa": st.session_state.get("followup_qa", []),
            "current_document": st.session_state.get("current_pdf_name", "")
        }

        # Add report data as JSON
        zip_file.writestr("report_data.json", json.dumps(report_data, indent=2))

        # Add analysis results
        for i, (result, ai_analysis) in enumerate(results):
            filename = result.get("filename", f"unknown_file_{i}")

            # Add the analysis result as JSON
            zip_file.writestr(
                f"analysis_results/{filename}_analysis.json",
                json.dumps(ai_analysis, indent=2)
            )

            # Add the annotated PDF if available
            if result.get("annotated_pdf"):
                try:
                    annotated_pdf_bytes = base64.b64decode(result["annotated_pdf"])
                    zip_file.writestr(
                        f"annotated_pdfs/{filename}_annotated.pdf",
                        annotated_pdf_bytes
                    )
                except Exception as e:
                    logger.error(f"Error adding annotated PDF for {filename} to report package: {e}")

            # Add the original document if available
            if "preprocessed_data" in st.session_state and filename in st.session_state.preprocessed_data:
                try:
                    # Try to get or restore original_bytes
                    original_bytes = restore_original_bytes_if_needed(filename)
                    if original_bytes:
                        zip_file.writestr(
                            f"original_documents/{filename}",
                            original_bytes
                        )
                except Exception as e:
                    logger.error(f"Error adding original document for {filename} to report package: {e}")

    # Reset buffer position and return the bytes
    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def export_to_word(exportable_results_list: List[Dict[str, Any]]) -> bytes:
    """
    Export analysis results to a Word document.

    Args:
        exportable_results_list: List of dictionaries containing analysis results

    Returns:
        bytes: The Word document as bytes
    """
    # Create a new Word document
    doc = Document()

    # Add title
    title = doc.add_heading("Document Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add report generation date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.italic = True

    def add_table_of_contents(target_doc: Document) -> None:
        """Insert a Word TOC field capturing heading levels 1-2."""
        try:
            target_doc.add_paragraph("Table of Contents", style="TOC Heading")
        except KeyError:
            target_doc.add_heading("Table of Contents", level=1)

        paragraph = target_doc.add_paragraph()
        begin_run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        begin_run._r.append(fld_begin)

        instr_run = paragraph.add_run()
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'
        instr_run._r.append(instr_text)

        separate_run = paragraph.add_run()
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        separate_run._r.append(fld_separate)

        placeholder_run = paragraph.add_run("(Press Ctrl+A then F9 to update the Table of Contents)")
        placeholder_run.italic = True

        end_run = paragraph.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        end_run._r.append(fld_end)

    add_table_of_contents(doc)

    # Add a page break after the title and TOC
    doc.add_page_break()

    method_priority = {
        'exact': 5,
        'exact_cleaned_search': 5,
        'special_case_quotes_handling': 3,
        'cross_page_fuzzy_match_part1': 2,
        'cross_page_fuzzy_match_part2': 2,
        'fuzzy': 2,
        'fuzzy_chunk_fallback_individual': 1,
        'fuzzy_chunk_fallback': 0
    }

    def pick_best_location(candidate_locs: List[Dict[str, Any]]) -> Dict[str, Any] | None:
        if not candidate_locs:
            return None

        def loc_key(loc: Dict[str, Any]) -> Tuple[int, float]:
            method = loc.get('method', '')
            score_val = loc.get('match_score', 0) or 0
            try:
                score_val = float(score_val)
            except Exception:
                score_val = 0.0
            return (method_priority.get(method, -1), score_val)

        return max(candidate_locs, key=loc_key)

    def normalize_candidate_locations(raw_locations: Any) -> List[Dict[str, Any]]:
        if isinstance(raw_locations, list):
            return [loc for loc in raw_locations if isinstance(loc, dict)]
        if isinstance(raw_locations, dict):
            if 'best_match' in raw_locations and isinstance(raw_locations['best_match'], dict):
                return [raw_locations['best_match']]
            return [raw_locations]
        return []

    def extract_clip_image(pdf_document: fitz.Document | None, location: Dict[str, Any], all_locations: List[Dict[str, Any]] = None) -> bytes | None:
        """
        Extract a clip image from the PDF. If multiple locations are provided on the same page,
        creates a bounding box that encompasses all of them (useful for exact match highlights
        that span multiple regions like clause number + full clause text).
        
        Args:
            pdf_document: The PDF document
            location: Primary location dict with page_num and rect
            all_locations: Optional list of all candidate locations for this phrase
            
        Returns:
            PNG image bytes or None if extraction fails
        """
        if pdf_document is None or not isinstance(location, dict):
            return None

        page_index = location.get("page_num")

        try:
            page_index = int(page_index)
        except Exception:
            return None

        try:
            page = pdf_document[page_index]
        except Exception as e:
            logger.error(f"Failed to access page {page_index} for citation snapshot: {e}")
            return None

        # Collect all rectangles on the same page
        rects_to_include = []
        
        # Add the primary location's rect
        rect_coords = location.get("rect")
        if isinstance(rect_coords, (list, tuple)) and len(rect_coords) == 4:
            rects_to_include.append(fitz.Rect(rect_coords))
        
        # If we have additional locations, check for same-page rects
        if all_locations and isinstance(all_locations, list):
            for loc in all_locations:
                if not isinstance(loc, dict):
                    continue
                    
                # Only include locations from the same page
                loc_page = loc.get("page_num")
                try:
                    loc_page = int(loc_page)
                except Exception:
                    continue
                    
                if loc_page == page_index:
                    loc_rect_coords = loc.get("rect")
                    if isinstance(loc_rect_coords, (list, tuple)) and len(loc_rect_coords) == 4:
                        loc_rect = fitz.Rect(loc_rect_coords)
                        # Only add if it's not a duplicate and not empty
                        if not loc_rect.is_empty and not any(loc_rect == r for r in rects_to_include):
                            rects_to_include.append(loc_rect)
        
        if not rects_to_include:
            return None
        
        # Create a combined bounding box that encompasses all rectangles
        clip_rect = rects_to_include[0]
        for rect in rects_to_include[1:]:
            clip_rect.include_rect(rect)
        
        if clip_rect.is_empty:
            return None

        # Pad the clip slightly so the excerpt has some context.
        try:
            page_rect = page.rect
            padding = 6.0
            clip_rect.x0 = max(page_rect.x0, clip_rect.x0 - padding)
            clip_rect.y0 = max(page_rect.y0, clip_rect.y0 - padding)
            clip_rect.x1 = min(page_rect.x1, clip_rect.x1 + padding)
            clip_rect.y1 = min(page_rect.y1, clip_rect.y1 + padding)
        except Exception:
            pass

        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), clip=clip_rect, alpha=False)
            return pix.tobytes("png")
        except Exception as e:
            logger.error(f"Failed to render citation snapshot: {e}")
            return None

    def add_markdown_runs(paragraph, text: str) -> None:
        """Render a limited subset of Markdown (bold/italic) into Word runs."""
        if not text:
            return

        lines = text.split('\n')
        for line_idx, line in enumerate(lines):
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', line)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**') and len(part) > 4:
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(part)
            if line_idx < len(lines) - 1:
                paragraph.add_run().add_break()

    def split_table_row(row_text: str) -> List[str]:
        row_text = row_text.strip()
        if row_text.startswith("|"):
            row_text = row_text[1:]
        if row_text.endswith("|"):
            row_text = row_text[:-1]
        return [cell.strip() for cell in row_text.split("|")]

    def is_table_separator(row_text: str) -> bool:
        stripped = row_text.strip()
        if not stripped or "|" not in stripped:
            return False
        stripped = stripped.strip("|")
        segments = stripped.split("|")
        if not segments:
            return False
        for segment in segments:
            segment = segment.strip()
            if not segment:
                return False
            if not re.fullmatch(r":?-{3,}:?", segment):
                return False
        return True

    def is_table_header(row_text: str) -> bool:
        stripped = row_text.strip()
        if "|" not in stripped:
            return False
        cells = split_table_row(stripped)
        return len(cells) >= 2

    def is_table_row(row_text: str) -> bool:
        stripped = row_text.strip()
        return bool(stripped) and "|" in stripped

    def infer_alignment(spec: str) -> str:
        spec = spec.strip()
        left = spec.startswith(":")
        right = spec.endswith(":")
        if left and right:
            return "center"
        if right:
            return "right"
        return "left"

    def parse_markdown_blocks(text: str) -> List[Dict[str, Any]]:
        """Parse limited Markdown into block structures (paragraphs, lists, tables)."""
        blocks: List[Dict[str, Any]] = []
        if not text:
            return blocks

        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            if is_table_header(line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                table_lines = [line, lines[i + 1]]
                i += 2
                while i < len(lines) and is_table_row(lines[i]):
                    if not lines[i].strip():
                        break
                    table_lines.append(lines[i])
                    i += 1
                blocks.append({"type": "table", "lines": table_lines})
                continue

            ul_match = re.match(r'^([\*\-\+])\s+(.+)$', stripped)
            if ul_match:
                items: List[str] = []
                while i < len(lines):
                    candidate = lines[i].strip()
                    bullet_match = re.match(r'^([\*\-\+])\s+(.+)$', candidate)
                    if not bullet_match:
                        break
                    items.append(bullet_match.group(2))
                    i += 1
                blocks.append({"type": "ul", "items": items})
                continue

            ol_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
            if ol_match:
                items: List[str] = []
                while i < len(lines):
                    candidate = lines[i].strip()
                    numbered_match = re.match(r'^(\d+)\.\s+(.+)$', candidate)
                    if not numbered_match:
                        break
                    items.append(numbered_match.group(2))
                    i += 1
                blocks.append({"type": "ol", "items": items})
                continue

            paragraph_lines = [lines[i]]
            i += 1
            while i < len(lines):
                next_line = lines[i]
                next_stripped = next_line.strip()
                if not next_stripped:
                    paragraph_lines.append(next_line)
                    i += 1
                    break
                if is_table_header(next_line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                    break
                if re.match(r'^([\*\-\+])\s+(.+)$', next_stripped) or re.match(r'^(\d+)\.\s+(.+)$', next_stripped):
                    break
                paragraph_lines.append(next_line)
                i += 1

            paragraph_text = '\n'.join(paragraph_lines).strip('\n')
            if paragraph_text:
                blocks.append({"type": "paragraph", "text": paragraph_text})

        return blocks

    def add_markdown_list(doc_obj: Document, items: List[str], ordered: bool) -> None:
        style = 'List Number' if ordered else 'List Bullet'
        for item in items:
            paragraph = doc_obj.add_paragraph(style=style)
            add_markdown_runs(paragraph, item)

    def add_markdown_table(doc_obj: Document, table_lines: List[str]) -> None:
        if len(table_lines) < 2:
            return

        header_cells = split_table_row(table_lines[0])
        if not header_cells:
            return

        alignment_specs = split_table_row(table_lines[1]) if len(table_lines) > 1 else []
        alignments = [infer_alignment(spec) for spec in alignment_specs]
        num_cols = len(header_cells)
        if len(alignments) < num_cols:
            alignments.extend(["left"] * (num_cols - len(alignments)))
        elif len(alignments) > num_cols:
            alignments = alignments[:num_cols]

        body_rows: List[List[str]] = []
        for raw_row in table_lines[2:]:
            if not is_table_row(raw_row):
                break
            cells = split_table_row(raw_row)
            if len(cells) < num_cols:
                cells.extend([""] * (num_cols - len(cells)))
            elif len(cells) > num_cols:
                cells = cells[:num_cols]
            body_rows.append(cells)

        table = doc_obj.add_table(rows=len(body_rows) + 1, cols=num_cols)
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            try:
                table.style = "Light Grid"
            except Exception:
                pass

        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        for col_idx, header_text in enumerate(header_cells):
            cell = table.cell(0, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            paragraph.alignment = alignment_map.get(alignments[col_idx], WD_ALIGN_PARAGRAPH.LEFT)
            add_markdown_runs(paragraph, header_text)
            for run in paragraph.runs:
                run.bold = True

        for row_idx, row_cells in enumerate(body_rows, start=1):
            for col_idx, cell_text in enumerate(row_cells):
                cell = table.cell(row_idx, col_idx)
                paragraph = cell.paragraphs[0]
                paragraph.text = ""
                paragraph.alignment = alignment_map.get(alignments[col_idx], WD_ALIGN_PARAGRAPH.LEFT)
                add_markdown_runs(paragraph, cell_text)

    def render_markdown_content(doc_obj: Document, text: str, initial_paragraph=None) -> None:
        blocks = parse_markdown_blocks(text)
        used_initial = False
        for block in blocks:
            block_type = block.get("type")
            if block_type == "paragraph":
                paragraph = initial_paragraph if initial_paragraph is not None and not used_initial else doc_obj.add_paragraph()
                add_markdown_runs(paragraph, block.get("text", ""))
            elif block_type == "ul":
                add_markdown_list(doc_obj, block.get("items", []), ordered=False)
            elif block_type == "ol":
                add_markdown_list(doc_obj, block.get("items", []), ordered=True)
            elif block_type == "table":
                add_markdown_table(doc_obj, block.get("lines", []))
            used_initial = True
            initial_paragraph = None

    # Process each file's results
    total_files = len(exportable_results_list)
    for file_index, file_result in enumerate(exportable_results_list):
        filename = file_result.get("filename", "Unknown File")
        analysis = file_result.get("analysis", {})
        phrase_details_map = file_result.get("phrase_details", {}) or {}
        phrase_locations_map = file_result.get("phrase_locations", {}) or {}
        annotated_pdf_b64 = file_result.get("annotated_pdf")
        section_to_prompt = file_result.get("section_to_prompt", {}) or {}
        annotated_pdf_doc = None

        if annotated_pdf_b64:
            try:
                annotated_pdf_doc = fitz.open(stream=base64.b64decode(annotated_pdf_b64), filetype="pdf")
            except Exception as pdf_err:
                logger.error(f"Unable to open annotated PDF for {filename}: {pdf_err}")
                annotated_pdf_doc = None

        try:
            # Add file heading
            doc.add_heading(f"Document: {filename}", 1)

            analysis_sections = list(analysis.get("analysis_sections", {}).items())
            total_sections = len(analysis_sections)

            # Process each analysis section
            for section_idx, (section_key, section_data) in enumerate(analysis_sections):
                if section_idx > 0:
                    doc.add_page_break()
                
                # Use original sub_prompt as section heading if available, otherwise use formatted section name
                original_prompt = section_to_prompt.get(section_key)
                if original_prompt:
                    doc.add_heading(original_prompt, 2)
                else:
                    section_name = section_key.replace("_", " ").title()
                    doc.add_heading(section_name, 2)

                # Add analysis text
                if section_data.get("Analysis"):
                    analysis_text = section_data.get("Analysis")
                    p = doc.add_paragraph()
                    label_run = p.add_run("Analysis: ")
                    label_run.bold = True
                    render_markdown_content(doc, analysis_text, initial_paragraph=p)

                # Add context if available
                if section_data.get("Context"):
                    p = doc.add_paragraph()
                    label_run = p.add_run("Context: ")
                    label_run.bold = True
                    add_markdown_runs(p, section_data.get("Context"))

                # Add supporting phrases
                supporting_phrases = section_data.get("Supporting_Phrases", [])
                if supporting_phrases and supporting_phrases != ["No relevant phrase found."]:
                    doc.add_heading("Supporting Citations", 3)

                    for phrase in supporting_phrases:
                        # Base verification info from flattened rows
                        is_verified = False
                        page_num_info = "Unknown"
                        score_info = "N/A"

                        try:
                            file_data = file_result.get("data", [])
                            if isinstance(file_data, list):
                                data_rows = [row for row in file_data if isinstance(row, dict) and row.get("Supporting Phrase") == phrase]
                            else:
                                data_rows = []
                        except Exception as e:
                            logger.error(f"Error getting data rows for phrase '{phrase}': {e}")
                            data_rows = []

                        if data_rows:
                            try:
                                data_row = data_rows[0]
                                verified_value = data_row.get("Verified")
                                if isinstance(verified_value, str):
                                    is_verified = verified_value.lower() == "yes"
                                elif isinstance(verified_value, bool):
                                    is_verified = verified_value

                                page_num_info = data_row.get("Page", "Unknown")
                                score_info = data_row.get("Match Score", "N/A")
                            except Exception as e:
                                logger.error(f"Error extracting verification info from data row: {e}")

                        phrase_info = phrase_details_map.get(phrase, {}) if isinstance(phrase_details_map, dict) else {}
                        candidate_locations = phrase_info.get("candidate_locations") if isinstance(phrase_info, dict) else None
                        if not candidate_locations:
                            candidate_locations = normalize_candidate_locations(phrase_locations_map.get(phrase))

                        best_location = phrase_info.get("best_location") if isinstance(phrase_info, dict) else None
                        if not best_location:
                            best_location = pick_best_location(candidate_locations or [])

                        if isinstance(phrase_info, dict) and "verified" in phrase_info:
                            is_verified = bool(phrase_info.get("verified"))

                        if best_location:
                            page_val = best_location.get("page_num")
                            if isinstance(page_val, int):
                                page_num_info = f"Page {page_val + 1}"
                            elif page_val is not None:
                                page_num_info = f"Page {page_val}"

                            match_val = best_location.get("match_score")
                            if match_val:
                                try:
                                    score_info = f"{float(match_val):.1f}%"
                                except Exception:
                                    score_info = str(match_val)

                        # Add the phrase with verification status
                        p = doc.add_paragraph()
                        if is_verified:
                            p.add_run("✓ ").bold = True
                            p.add_run(phrase)
                            details_run = p.add_run(f" (Pg: {page_num_info}, Score: {score_info})")
                            details_run.italic = True
                            details_run.font.size = Pt(9)
                        else:
                            p.add_run("❓ ").bold = True
                            p.add_run(phrase)
                            details_run = p.add_run(" (Not verified in document)")
                            details_run.italic = True
                            details_run.font.size = Pt(9)

                        # Attach a citation snapshot when possible.
                        # Pass all candidate locations so the image can capture multiple highlights on the same page
                        # (e.g., clause number "5.3" + full clause text as separate highlights)
                        snapshot_bytes = extract_clip_image(annotated_pdf_doc, best_location, candidate_locations) if best_location else None
                        if snapshot_bytes:
                            try:
                                img_stream = BytesIO(snapshot_bytes)
                                img_paragraph = doc.add_paragraph()
                                img_run = img_paragraph.add_run()
                                img_run.add_picture(img_stream, width=Inches(6.0))
                                caption = doc.add_paragraph("Excerpt from annotated PDF")
                                try:
                                    caption.style = "Caption"
                                except KeyError:
                                    pass
                            except Exception as img_err:
                                logger.error(f"Error attaching citation snapshot for phrase '{phrase}': {img_err}")

        finally:
            if annotated_pdf_doc:
                try:
                    annotated_pdf_doc.close()
                except Exception:
                    pass

        # Ensure separation between files when more remain
        if file_index < total_files - 1:
            doc.add_page_break()

    # Add Follow-up Q&A section if any exist
    followup_qa = st.session_state.get("followup_qa", [])
    if followup_qa:
        doc.add_heading("Follow-up Questions & Answers", 1)

        for i, qa_pair in enumerate(followup_qa):
            # Add question
            doc.add_heading(f"Question {i+1}", 2)
            q_paragraph = doc.add_paragraph()
            q_paragraph.add_run("Q: ").bold = True
            q_paragraph.add_run(qa_pair.get("question", ""))

            # Add answer
            a_paragraph = doc.add_paragraph()
            a_paragraph.add_run("A: ").bold = True
            a_paragraph.add_run(qa_pair.get("answer", ""))

            # Add timestamp if available
            if qa_pair.get("timestamp"):
                timestamp_paragraph = doc.add_paragraph()
                timestamp_run = timestamp_paragraph.add_run(f"Asked on: {qa_pair['timestamp']}")
                timestamp_run.italic = True
                timestamp_run.font.size = Pt(9)

            # Add citations if available
            citation_details = qa_pair.get("citation_details", [])
            if citation_details:
                doc.add_heading("Citations", 3)
                for citation in citation_details:
                    cite_paragraph = doc.add_paragraph()
                    cite_paragraph.style = 'List Bullet'
                    cite_paragraph.add_run(f"[{citation.get('number', '')}] ")
                    cite_paragraph.add_run(f"{citation.get('filename', 'Unknown')}, ")
                    cite_paragraph.add_run(f"Page {citation.get('page_num', 'Unknown')}")

            # Add separator between Q&A pairs
            if i < len(followup_qa) - 1:
                doc.add_paragraph("---")

    # Save the document to a BytesIO object
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    # Return the document as bytes
    return docx_buffer.getvalue()

